"""
Document chunker – extracts text from PDF and MS Word files, then splits it
into overlapping chunks suitable for AI model ingestion.

Supported formats
-----------------
- PDF  (.pdf)  via ``pypdf``
- Word (.docx) via ``python-docx``
- Word (.doc)  falls back to plain-text extraction via ``python-docx`` where
               possible; legacy .doc files may require LibreOffice conversion.
"""

from __future__ import annotations

import io
import logging
import os
from dataclasses import dataclass, field
from typing import List

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Optional dependency guards
# ---------------------------------------------------------------------------
try:
    from pypdf import PdfReader
    _PYPDF_AVAILABLE = True
except ImportError:  # pragma: no cover
    _PYPDF_AVAILABLE = False
    logger.warning("pypdf not installed – PDF extraction unavailable")

try:
    from docx import Document as DocxDocument
    _DOCX_AVAILABLE = True
except ImportError:  # pragma: no cover
    _DOCX_AVAILABLE = False
    logger.warning("python-docx not installed – Word extraction unavailable")


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------

@dataclass
class TextChunk:
    """A single text chunk produced by the :class:`Chunker`."""

    patient_id: str
    source_file: str
    chunk_index: int
    total_chunks: int
    text: str
    page_numbers: List[int] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Extraction helpers
# ---------------------------------------------------------------------------

def _extract_text_pdf(content: bytes) -> List[tuple[str, List[int]]]:
    """
    Extract text from a PDF byte string.

    Returns a list of ``(paragraph_text, [page_number])`` tuples.
    """
    if not _PYPDF_AVAILABLE:
        raise ImportError("pypdf is required for PDF extraction. Install it with: pip install pypdf")

    reader = PdfReader(io.BytesIO(content))
    paragraphs: List[tuple[str, List[int]]] = []
    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        for para in text.split("\n"):
            para = para.strip()
            if para:
                paragraphs.append((para, [page_num]))
    return paragraphs


def _extract_text_docx(content: bytes) -> List[tuple[str, List[int]]]:
    """
    Extract text from a .docx byte string.

    Returns a list of ``(paragraph_text, [])`` tuples (Word docs have no
    intrinsic page-number metadata at paragraph level).
    """
    if not _DOCX_AVAILABLE:
        raise ImportError("python-docx is required for Word extraction. Install it with: pip install python-docx")

    doc = DocxDocument(io.BytesIO(content))
    paragraphs: List[tuple[str, List[int]]] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append((text, []))
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append((row_text, []))
    return paragraphs


def extract_text(content: bytes, filename: str) -> List[tuple[str, List[int]]]:
    """Dispatch to the correct extractor based on the file extension."""
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        return _extract_text_pdf(content)
    if ext in (".docx", ".doc"):
        return _extract_text_docx(content)
    raise ValueError(f"Unsupported file type: {ext!r}")


# ---------------------------------------------------------------------------
# Chunker
# ---------------------------------------------------------------------------

class Chunker:
    """
    Splits a document's extracted text into overlapping chunks.

    Parameters
    ----------
    chunk_size : int
        Maximum number of *characters* per chunk (default 2 000).
    overlap : int
        Number of characters to repeat at the start of the next chunk
        (default 200).  Overlap ensures context is not lost at boundaries.
    min_chunk_size : int
        Chunks smaller than this number of characters are merged with the
        previous chunk rather than emitted as standalone chunks (default 100).
    """

    def __init__(
        self,
        chunk_size: int = 2000,
        overlap: int = 200,
        min_chunk_size: int = 100,
    ) -> None:
        if chunk_size <= 0:
            raise ValueError("chunk_size must be positive")
        if overlap < 0:
            raise ValueError("overlap must be non-negative")
        if overlap >= chunk_size:
            raise ValueError("overlap must be less than chunk_size")
        self.chunk_size = chunk_size
        self.overlap = overlap
        self.min_chunk_size = min_chunk_size

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def chunk_document(
        self,
        content: bytes,
        filename: str,
        patient_id: str,
    ) -> List[TextChunk]:
        """
        Extract text from *content* and split it into overlapping chunks.

        Parameters
        ----------
        content : bytes
            Raw file bytes.
        filename : str
            Original filename (used to determine the parser and stored as
            metadata on each chunk).
        patient_id : str
            Patient identifier stored on each chunk.

        Returns
        -------
        list of :class:`TextChunk`
        """
        paragraphs = extract_text(content, filename)
        if not paragraphs:
            logger.warning("No text extracted from %s", filename)
            return []

        raw_chunks = self._build_raw_chunks(paragraphs)
        chunks: List[TextChunk] = []
        total = len(raw_chunks)
        for idx, (text, pages) in enumerate(raw_chunks):
            chunks.append(
                TextChunk(
                    patient_id=patient_id,
                    source_file=filename,
                    chunk_index=idx,
                    total_chunks=total,
                    text=text,
                    page_numbers=pages,
                )
            )
        return chunks

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    def _build_raw_chunks(
        self,
        paragraphs: List[tuple[str, List[int]]],
    ) -> List[tuple[str, List[int]]]:
        """
        Combine paragraphs into fixed-size overlapping text windows.

        Returns a list of ``(chunk_text, list_of_page_numbers)`` tuples.
        """
        # First, flatten all paragraphs into a single character stream while
        # tracking page boundaries.
        full_text = ""
        page_map: List[int] = []  # page_map[i] = page number for character i
        for para, pages in paragraphs:
            page_num = pages[0] if pages else 0
            full_text += para + "\n"
            page_map.extend([page_num] * (len(para) + 1))

        if not full_text.strip():
            return []

        raw_chunks: List[tuple[str, List[int]]] = []
        start = 0
        text_len = len(full_text)

        while start < text_len:
            end = min(start + self.chunk_size, text_len)
            chunk_text = full_text[start:end].strip()

            if len(chunk_text) < self.min_chunk_size and raw_chunks:
                # Merge tiny trailing chunk into the previous one
                prev_text, prev_pages = raw_chunks[-1]
                raw_chunks[-1] = (prev_text + " " + chunk_text, prev_pages)
            else:
                chunk_pages = sorted(set(page_map[start:end]))
                raw_chunks.append((chunk_text, chunk_pages))

            if end == text_len:
                break
            start = end - self.overlap

        return raw_chunks
